
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_detailed_report():
    doc = Document()
    
    # --- STYLES ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    h1_style = doc.styles['Heading 1']
    h1_style.font.color.rgb = RGBColor(46, 116, 181)
    
    h2_style = doc.styles['Heading 2']
    h2_style.font.color.rgb = RGBColor(68, 114, 196)

    # --- TITLE PAGE ---
    title = doc.add_heading('Luxe Stay – Hotel Management System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Technical Report & Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    doc.add_paragraph('\n\n')
    doc.add_paragraph('Developed by: [Your Name]').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Semester 2 | PF Programming').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 1. EXECUTIVE SUMMARY ---
    doc.add_heading('1. Executive Summary', level=1)
    p = doc.add_paragraph(
        "The 'Luxe Stay' Hotel Management System is a versatile, multi-interface software solution tailored for "
        "the hospitality industry. It bridges the gap between traditional desktop management and modern web accessibility "
        "by offering three distinct user interfaces: a robust Command Line Interface (CLI), a Native Desktop GUI, and a "
        "responsive Web Application. All interfaces share a centralized backend and data layer, ensuring real-time synchronization."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- 2. SYSTEM ARCHITECTURE ---
    doc.add_heading('2. System Architecture', level=1)
    doc.add_paragraph("The system is built on a modular architecture using Python 3.13.")
    
    # Modules Table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = "Module / File"
    hdr[1].text = "Functionality"
    
    components = [
        ("main.py", "Core Backend Logic & Console Interface. Handles data loading, saving (Local/Cloud), and business rules."),
        ("app.py", "Flask Web Server. Exposes REST API endpoints for the web frontend to communicate with the backend."),
        ("server.py", "Production Entry Point. Uses 'Waitress' WSGI server to host the web app reliably on Windows."),
        ("gui_app.py", "Native Desktop Application. Built with CustomTkinter to provide a modern, dark-themed GUI similar to the web app."),
        ("index.html / script.js", "Web Frontend. Dynamic Single Page Application (SPA) for browser-based access."),
        ("Google Apps Script", "Cloud Database Connector. Acts as a bridge between the Python backend and Google Sheets.")
    ]
    
    for mod, desc in components:
        row = table.add_row().cells
        row[0].text = mod
        row[1].text = desc

    # --- 3. CONSOLE INTERFACE (CLI) ---
    doc.add_heading('3. Console Interface (main.py)', level=1)
    doc.add_paragraph(
        "The Console Interface serves as the foundational layer of the application. It runs directly in the terminal "
        "and provides a fail-safe method for system administration."
    )
    
    doc.add_heading('Features:', level=2)
    cli_features = [
        "Robust Error Handling: Validates all user inputs (e.g., preventing negative prices, checking room duplicates).",
        "Role-Based Menus: The 'Administrator Dashboard' offers the full suite of tools, while other roles are restricted.",
        "Data Persistence Engine: The CLI contains the core logic for reading/writing to text files (data/rooms.txt, etc.) and synchronizing with the cloud.",
        "Reporting: Generates text-based tables for Financial Reports, Staff Lists, and Room Inventory."
    ]
    for feat in cli_features:
        doc.add_paragraph(f"• {feat}", style='List Bullet')

    doc.add_paragraph().add_run("[INSERT CONSOLE SCREENSHOT HERE]").font.color.rgb = RGBColor(255, 0, 0)


    # --- 4. WEB INTERFACE (Frontend) ---
    doc.add_heading('4. Web Interface', level=1)
    doc.add_paragraph(
        "Top-tier user experience design branding known as the 'Royal Suite' theme. "
        "It uses a luxurious Black & Gold color palette to convey premium service."
    )
    
    doc.add_heading('Technical Details:', level=2)
    doc.add_paragraph(
        "• Technology: HTML5, CSS3 (Flexbox/Grid), Vanilla JavaScript (ES6+).\n"
        "• Communication: Uses `fetch` API to talk to the Flask backend endpoints (e.g., /api/login, /api/stats).\n"
        "• Responsiveness: Fully responsive layout that adapts to different screen sizes."
    )

    doc.add_heading('Key Web Features:', level=3)
    web_feats = [
        ("Dynamic Dashboard", "Widgets showing Real-time Revenue, Occupancy, and Active Staff."),
        ("Live Navigation", "Sidebar navigation highlights the active tab."),
        ("Guest Management", "Receptionists can Book Rooms (filtering available ones) and Check Out guests with one click."),
        ("Admin Controls", "Admins can Delete Records (Rooms/Staff) and Update Settings directly from the UI.")
    ]
    for title, desc in web_feats:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    if os.path.exists("web_screenshot.png"):
        doc.add_picture("web_screenshot.png", width=Inches(5.5))

    # --- 5. DESKTOP GUI ---
    doc.add_heading('5. Desktop GUI (gui_app.py)', level=1)
    doc.add_paragraph(
        "The Desktop GUI brings the web's aesthetics to a native Windows window using the 'CustomTkinter' library."
    )
    
    gui_feats = [
        "Visual Consistency: Matches the Web App's Dark/Gold theme pixel-for-pixel.",
        "Native Performance: Runs locally without a browser.",
        "Modal Dialogs: Custom pop-up windows for data entry (Add Room, Assign Task).",
        "Scrollable Frames: Handled large datasets (Guest Lists) smoothly."
    ]
    for feat in gui_feats:
        doc.add_paragraph(f"• {feat}", style='List Bullet')

    doc.add_paragraph().add_run("[INSERT GUI SCREENSHOT HERE]").font.color.rgb = RGBColor(255, 0, 0)

    # --- 6. EXTENDED FEATURE LIST ---
    doc.add_heading('6. Feature Detail Breakdown', level=1)
    
    doc.add_heading('A. Role-Based Access Control (RBAC)', level=2)
    doc.add_paragraph("The system secures data by assigning users one of four roles:")
    doc.add_paragraph(
        "1. Admin: Full system access (Manage Rooms, Staff, Finance, Settings).\n"
        "2. Manager: Operational access (Staff Attendance, Tasks, Reports).\n"
        "3. Receptionist: Guest facing (Bookings, Check-ins, Check-outs).\n"
        "4. Worker: Task execution (View personal tasks, Clock In/Out)."
    )

    doc.add_heading('B. Cloud Synchronization', level=2)
    doc.add_paragraph(
        "A unique feature of this system is its Hybrid Database. While it operates on local text files for speed, "
        "it is capable of syncing with Google Sheets via a custom Google Apps Script. This allows owners to check "
        "hotel stats remotely from their phone via the Google Sheets app."
    )

    doc.add_heading('C. Task Management', level=2)
    doc.add_paragraph(
        "Admins can assign tasks to specific workers. Workers see these tasks on their personal dashboard and "
        "can mark them as 'Completed'. This replaces paper-based to-do lists."
    )

    # --- 7. CONCLUSION ---
    doc.add_heading('7. Conclusion', level=1)
    doc.add_paragraph(
        "This project demonstrates a full-stack capability, handling low-level file I/O in Python while delivering "
        "high-level user experiences via Web and GUI. The consistent design language and robust feature set make it "
        "a production-ready prototype for luxury hotel management."
    )

    # Save
    filename = 'Detailed_Project_Report.docx'
    doc.save(filename)
    print(f"Report generated: {filename}")

if __name__ == "__main__":
    create_detailed_report()
